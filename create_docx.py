import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_docx():
    doc = Document()
    
    # Page Setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Title / Heading 1
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    run1 = h1.add_run("6.3 Shayan Naghibi, Backend")
    run1.font.name = 'Arial'
    run1.font.size = Pt(20)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    def add_h2(text):
        h2 = doc.add_heading(level=2)
        h2.paragraph_format.space_before = Pt(16)
        h2.paragraph_format.space_after = Pt(6)
        run = h2.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        return h2

    def add_h3(text):
        h3 = doc.add_heading(level=3)
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(4)
        run = h3.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x36, 0x60, 0x92)
        return h3

    def add_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
        p.add_run(text)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
        p.add_run(text)
        return p

    # --- 6.3.1 ---
    add_h2("6.3.1 What I personally contributed")

    add_h3("Early authentication and registration core")
    add_p("During Sprint 1 in May 2026, I worked on the initial backend registration and user profile queries in Backend/main.py. I implemented the duplicate email validation check on sign-up (POST /auth/register), the user login endpoint (POST /auth/login), the user profile lookup route (GET /api/users/{user_id}), and the environment path resolution logic in commit 1198ba9. This ensured that registration errors were caught gracefully and that environment variables loaded reliably regardless of the working directory from which the server was launched.")

    add_h3("The file management subsystem")
    add_p("Later in the project, I built the complete File Upload and Download Subsystem. I designed and implemented an independent backend router (Backend/files.py) mounted cleanly into FastAPI, supporting multipart file uploads (POST /api/files/upload), file metadata listing (GET /api/files/list), binary file downloading (GET /api/files/download/{file_id}), and role-restricted file deletion (DELETE /api/files/{file_id}). On the frontend, I created the reusable FileManager.jsx component, added API handlers in api.js, and integrated a new \"Files\" section across the Student, Teacher, and Admin dashboards in commit 38b2dfa.")

    # --- 6.3.2 ---
    add_h2("6.3.2 Technical tasks I completed")

    add_h3("Commits submitted")
    table1_data = [
        ["Commit", "Date", "What I built"],
        ["1198ba9", "28 May 2026", "Added duplicate email validation on sign-up, .env path resolution, login query (POST /auth/login), and profile lookup (GET /api/users/{user_id})."],
        ["38b2dfa", "12 Aug 2026", "Implemented file upload, download, list, and delete endpoints in files.py, created FileManager.jsx, and integrated the Files tab into all dashboards."]
    ]
    t1 = doc.add_table(rows=len(table1_data), cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table1_data):
        for c_idx, val in enumerate(row):
            cell = t1.cell(r_idx, c_idx)
            cell.text = val
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            if r_idx == 0:
                set_cell_background(cell, "1F497D")
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F2F5F9")
                else:
                    set_cell_background(cell, "FFFFFF")

    add_h3("Endpoints I implemented")
    table2_data = [
        ["Endpoint", "Method", "Auth", "Purpose"],
        ["/auth/register", "POST", "None", "Sign up user; returns explicit error if email is already registered."],
        ["/auth/login", "POST", "None", "Authenticates user credentials against Supabase Auth."],
        ["/api/users/{user_id}", "GET", "Token", "Looks up user account details by UUID."],
        ["/api/files/upload", "POST", "Token", "Uploads course material with optional description to server storage."],
        ["/api/files/list", "GET", "Token", "Lists all uploaded course materials with metadata."],
        ["/api/files/download/{file_id}", "GET", "Token", "Streams binary file download to client."],
        ["/api/files/{file_id}", "DELETE", "Token", "Deletes a file (restricted to uploader or admin)."]
    ]
    t2 = doc.add_table(rows=len(table2_data), cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table2_data):
        for c_idx, val in enumerate(row):
            cell = t2.cell(r_idx, c_idx)
            cell.text = val
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            if r_idx == 0:
                set_cell_background(cell, "1F497D")
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F2F5F9")
                else:
                    set_cell_background(cell, "FFFFFF")

    add_h3("Frontend components & tooling")
    add_bullet("Created frontend/src/components/FileManager.jsx featuring drag-and-drop file upload, custom descriptions, file type icon badges, size formatters, real-time search filtering, and action buttons.")
    add_bullet("Added API helper functions in frontend/src/api.js: uploadFileApi(), fetchFilesApi(), downloadFileApi(), and deleteFileApi().")
    add_bullet("Added SVG icon to SidebarIcons.jsx and updated navigation menus in DashboardStudentPage.jsx, DashboardTeacherPage.jsx, and DashboardAdminPage.jsx.")
    add_bullet("Resolved virtualenv installation issues on Windows PowerShell by installing python-multipart for FastAPI form handling and lucide-react for frontend build resolution.")

    # --- 6.3.3 ---
    add_h2("6.3.3 Design decisions I influenced")

    add_h3("1. Explicit duplicate email error responses")
    add_p("In commit 1198ba9, I caught generic exceptions from Supabase Auth during registration and converted duplicate email errors into explicit validation messages (\"Validation Failed: This email is already registered.\"). This prevented raw backend stack traces from reaching the client and provided readable feedback to the frontend team.")

    add_h3("2. Absolute directory resolution for environment variables")
    add_p("To fix environment loading errors when executing python main.py from different working directories, I updated main.py to resolve .env relative to the script file using os.path.dirname(os.path.abspath(__file__)). This ensured .env is always located reliably before initializing the Supabase client.")

    add_h3("3. Standalone router for file management (files.py)")
    add_p("Rather than adding more code into Backend/main.py where multiple team members were committing, I isolated file management into a dedicated files_router in Backend/files.py and attached it via app.include_router(files_router). This kept the codebase modular and eliminated merge conflicts.")

    add_h3("4. Dual storage strategy with local fallback")
    add_p("To ensure file upload and download worked out-of-the-box in local development without requiring manual Supabase Storage bucket setup, I designed Backend/files.py to save uploaded files to Backend/uploads/ with metadata stored in metadata.json. This provided instant plug-and-play functionality while remaining compatible with cloud storage migration.")

    add_h3("5. Role-restricted file deletion")
    add_p("I enforced authorization checks on file deletion (DELETE /api/files/{file_id}), verifying that the user requesting deletion is either the original uploader or an administrator.")

    # --- 6.3.4 ---
    add_h2("6.3.4 Problems I solved")

    add_h3("Duplicate email registration crashes")
    add_bullet("Attempting to sign up with an email that was already registered triggered unhandled exceptions or generic 500 responses.", "The problem: ")
    add_bullet("Manual API testing during initial authentication setup.", "How I found it: ")
    add_bullet("Wrapped registration in a try/except block that checks for \"already registered\" or \"already exists\" in the Supabase error payload and returns a clean status: \"Error\" message.", "The fix, in 1198ba9: ")
    add_bullet("A small string inspection step during sign-up error handling, yielding significantly clearer user feedback.", "What it cost: ")

    add_h3("Environment variable loading failures")
    add_bullet("Running python main.py from outside the Backend/ directory failed because load_dotenv() looked in the current shell directory rather than the backend folder.", "The problem: ")
    add_bullet("Testing backend startup from different root folders.", "How I found it: ")
    add_bullet("Derived dotenv_path = os.path.join(current_dir, \".env\") and added explicit validation raising ValueError if key variables are missing.", "The fix, in 1198ba9: ")

    add_h3("FastAPI multipart form crash")
    add_bullet("Calling POST /api/files/upload raised RuntimeError: Form data requires \"python-multipart\" to be installed.", "The problem: ")
    add_bullet("Testing file upload route invocation in FastAPI.", "How I found it: ")
    add_bullet("Installed python-multipart into Backend/venv and added it to Backend/requirements.txt.", "The fix: ")

    add_h3("Vite production build failures")
    add_bullet("Running npm run build failed due to missing lucide-react dependency and a duplicate body key in RegisterPage.jsx.", "The problem: ")
    add_bullet("Validating frontend production compilation prior to merge.", "How I found it: ")
    add_bullet("Installed lucide-react and cleaned RegisterPage.jsx, achieving a 100% clean Vite build with 0 errors.", "The fix: ")

    # --- 6.3.5 ---
    add_h2("6.3.5 What I learned")

    add_h3("FastAPI multipart streaming and file handling")
    add_p("I gained practical experience using FastAPI's UploadFile and FileResponse classes to process multipart form uploads and stream binary file downloads. I learned how to handle filename sanitization using UUIDs to prevent file overwrites on disk, and how to construct clean metadata JSON structures for client consumption.")

    add_h3("Resilient environment configuration and dependency management")
    add_p("Working on Windows PowerShell taught me how to manage execution policies, set up isolated Python virtual environments, and resolve environment path issues using os.path.abspath. I also learned how to diagnose missing package requirements quickly (such as python-multipart and lucide-react) and update dependency files (requirements.txt and package.json).")

    add_h3("Non-disruptive software architecture")
    add_p("By building files.py as a standalone router and FileManager.jsx as a self-contained component, I learned how to extend a shared codebase without altering existing team members' code. This modular approach prevented regressions and simplified Git merging.")

    # --- 6.3.6 ---
    add_h2("6.3.6 How I grew professionally")

    add_h3("Shifting from isolated functions to full-stack feature ownership")
    add_p("At the start of the project, I worked on isolated backend snippets. By the end, I took full ownership of an end-to-end subsystem—from designing the REST API routes and file storage mechanics in FastAPI to building the interactive drag-and-drop React interface in Vite.")

    add_h3("Designing defensive, fallback-capable software")
    add_p("Building the file management system taught me to anticipate missing external services. By creating a local storage fallback (Backend/uploads/), I ensured that the feature remained functional even if cloud storage buckets were unconfigured in local development environments.")

    add_h3("Collaborating effectively in a multi-developer Git repository")
    add_p("Working on Shayan-branch, merging updates from main, and resolving branch drift taught me the importance of clean Git workflows. I learned to isolate my work into feature branches, keep commits focused, and write code that integrates seamlessly into a shared team repository.")

    # Save to file
    out_path_1 = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Section_6_3_Shayan_Naghibi.docx")
    doc.save(out_path_1)
    print(f"Successfully generated docx file at: {out_path_1}")

if __name__ == "__main__":
    generate_docx()
